# app/services/export_service.py

from PIL import Image
from io import BytesIO
import os
import logging
from datetime import datetime
from typing import Optional, List, Dict, Any
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from reportlab.lib import colors
from reportlab.lib.pagesizes import A4, LETTER, LEGAL
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, Image as RLImage
from werkzeug.exceptions import BadRequest

logger = logging.getLogger(__name__)

DEFAULT_EXPORT_PARAMS = {
    'page_size': 'A4',
    'margin_top': 1.0,
    'margin_bottom': 1.0,
    'margin_left': 1.0,
    'margin_right': 1.0,
    'line_spacing': 1.15,
    'font_size': 12,
    'logo_path': None,
    'signatures': [
        {
            'title': 'Completed by',
            'name': '',
            'date': True
        },
        {
            'title': 'Reviewed by',
            'name': '',
            'date': True
        }
    ],
    'signature_spacing': {
        'before_section': 20,  # Space before signatures section
        'between_signatures': 10,  # Space between signatures
        'after_date': 5,  # Space after date field
        'after_section': 20  # Space after signatures section
    }
}

class ExportService:
    def __init__(self):
        self.supported_formats = ['PDF', 'DOCX']
        self.page_sizes = {
            'A4': A4,
            'LETTER': LETTER,
            'LEGAL': LEGAL
        }

    def _validate_form_data(self, form_data: Dict[str, Any]) -> None:
        """Validate form data structure before export"""
        required_fields = ['title', 'created_by', 'questions']
        for field in required_fields:
            if field not in form_data:
                raise ValueError(f"Missing required field: {field}")
            
        if not isinstance(form_data['questions'], list):
            raise ValueError("Questions must be a list")
        
        for question in form_data['questions']:
            if 'text' not in question or 'type' not in question:
                raise ValueError("Each question must have 'text' and 'type' fields")

    def _add_signatures_pdf(self, story: List, signatures: List[Dict], styles: Dict, spacing_params: Dict) -> None:
        """Add signature section to PDF document with customizable spacing"""
        # Space before signatures section
        story.append(Spacer(1, spacing_params.get('before_section', 20)))
        story.append(Paragraph("Signatures:", styles['Heading2']))
        
        for i, sig in enumerate(signatures):
            # Add signature block
            signature_block = []
            
            # Title and signature line
            signature_block.append(
                Paragraph(f"{sig['title']}: {'_' * 40}", styles['FormField'])
            )
            
            # Pre-filled name if provided
            if sig.get('name'):
                signature_block.append(
                    Paragraph(f"Name: {sig['name']}", styles['FormField'])
                )
            else:
                signature_block.append(
                    Paragraph("Name: _________________________________", styles['FormField'])
                )
            
            # Date field if requested
            if sig.get('date', True):
                signature_block.append(
                    Paragraph("Date: ____________________", styles['FormField'])
                )
                # Space after date field
                signature_block.append(Spacer(1, spacing_params.get('after_date', 5)))
            
            story.append(Table([[cell] for cell in signature_block],
                            spaceBefore=0,
                            spaceAfter=spacing_params.get('between_signatures', 10)))
            
        # Space after signatures section
        story.append(Spacer(1, spacing_params.get('after_section', 20)))

    def _add_signatures_docx(self, doc: Document, signatures: List[Dict], spacing_params: Dict) -> None:
        """Add signature section to DOCX document with customizable spacing"""
        doc.add_paragraph().add_run().add_break()
        doc.add_heading('Signatures:', level=1)
        
        # Space before first signature
        doc.add_paragraph().add_run().add_break()
        
        for i, sig in enumerate(signatures):
            # Add signature block
            p = doc.add_paragraph()
            p.style.paragraph_format.space_before = Pt(spacing_params.get('before_section', 20))
            p.add_run(f"{sig['title']}: ").bold = True
            p.add_run("_" * 40)
            
            # Pre-filled name if provided
            p = doc.add_paragraph()
            p.style.paragraph_format.space_before = Pt(spacing_params.get('after_date', 5))
            p.add_run("Name: ").bold = True
            if sig.get('name'):
                p.add_run(sig['name'])
            else:
                p.add_run("_" * 40)
            
            # Date field if requested
            if sig.get('date', True):
                p = doc.add_paragraph()
                p.style.paragraph_format.space_before = Pt(spacing_params.get('after_date', 5))
                p.add_run("Date: ").bold = True
                p.add_run("_" * 20)
            
            # Space between signatures
            if i < len(signatures) - 1:
                p = doc.add_paragraph()
                p.style.paragraph_format.space_after = Pt(spacing_params.get('between_signatures', 10))

    def _add_logo_pdf(self, story: List, logo_path: str, width: float = 2.0) -> None:
        """Add logo to PDF document"""
        try:
            if os.path.exists(logo_path):
                img = Image.open(logo_path)
                aspect = img.height / img.width
                img_width = width * inch
                img_height = img_width * aspect
                story.append(RLImage(logo_path, width=img_width, height=img_height))
                story.append(Spacer(1, 20))
        except Exception as e:
            logger.warning(f"Could not add logo to PDF: {str(e)}")

    def _add_logo_docx(self, doc: Document, logo_path: str, width: float = 2.0) -> None:
        """Add logo to DOCX document"""
        try:
            if os.path.exists(logo_path):
                if logo_path.lower().endswith(('.png', '.jpg', '.jpeg')):
                    doc.add_picture(logo_path, width=Inches(width))
                    doc.add_paragraph()  # Add spacing after logo
        except Exception as e:
            logger.warning(f"Could not add logo to DOCX: {str(e)}")

    def export_as_pdf(self, form_data: Dict[str, Any], format_params: Dict[str, Any]) -> bytes:
        """Export form as fillable PDF with custom formatting"""
        try:
            self._validate_form_data(form_data)

            buffer = BytesIO()
            page_size = self.page_sizes[format_params['page_size']]
            
            # Create document
            doc = SimpleDocTemplate(
                buffer,
                pagesize=page_size,
                rightMargin=format_params['margin_right'] * inch,
                leftMargin=format_params['margin_left'] * inch,
                topMargin=format_params['margin_top'] * inch,
                bottomMargin=format_params['margin_bottom'] * inch
            )

            # Prepare styles
            styles = getSampleStyleSheet()
            base_font_size = format_params['font_size']
            line_spacing = format_params['line_spacing']

            # Custom styles
            styles.add(ParagraphStyle(
                name='FormTitle',
                parent=styles['Heading1'],
                fontSize=base_font_size + 4,
                spaceAfter=20,
                alignment=1,
                leading=base_font_size * line_spacing
            ))

            styles.add(ParagraphStyle(
                name='FormField',
                parent=styles['Normal'],
                fontSize=base_font_size,
                spaceAfter=15,
                leading=base_font_size * line_spacing
            ))

            styles.add(ParagraphStyle(
                name='Answer',
                parent=styles['Normal'],
                fontSize=base_font_size,
                leftIndent=20,
                spaceAfter=10,
                leading=base_font_size * line_spacing
            ))

            # Build content
            story = []

            # Add logo if provided
            if format_params.get('logo_path'):
                self._add_logo_pdf(story, format_params['logo_path'])

            # Title and description
            story.append(Paragraph(form_data['title'], styles['FormTitle']))
            if form_data.get('description'):
                story.append(Paragraph(form_data['description'], styles['Normal']))
            story.append(Spacer(1, 20))

            # Form Information
            story.append(Paragraph("Form Information:", styles['Heading2']))
            story.append(Paragraph(
                f"Environment: {form_data['created_by']['environment']['name']}", 
                styles['Normal']
            ))
            story.append(Paragraph(
                f"Created by: {form_data['created_by']['fullname']}", 
                styles['Normal']
            ))
            story.append(Spacer(1, 20))

            # Response Information
            story.append(Paragraph("Response Information:", styles['Heading2']))
            story.append(Paragraph("Name: _________________________________", styles['FormField']))
            story.append(Paragraph("Date: _________________________________", styles['FormField']))
            story.append(Spacer(1, 20))

            # Questions
            for i, question in enumerate(form_data['questions'], 1):
                # Question text
                story.append(Paragraph(f"{i}. {question['text']}", styles['FormField']))
                
                # Handle different question types
                if question['type'] == 'text':
                    story.append(Paragraph("_________________________________", styles['Answer']))
                
                elif question['type'] in ['checkbox', 'multiple_choices']:
                    if question.get('possible_answers'):
                        for answer in question['possible_answers']:
                            story.append(Paragraph(f"□ {answer['value']}", styles['Answer']))
                    else:
                        story.append(Paragraph("□ Yes    □ No", styles['Answer']))

                # Add remarks if present
                if question.get('remarks'):
                    story.append(Paragraph(f"Remarks: {question['remarks']}", styles['Answer']))

                story.append(Spacer(1, 10))

            # Add signatures
            if 'signatures' in format_params:
                self._add_signatures_pdf(
                    story,
                    format_params['signatures'],
                    styles,
                    format_params.get('signature_spacing', DEFAULT_EXPORT_PARAMS['signature_spacing'])
                )

            # Build PDF
            doc.build(story)
            return buffer.getvalue()

        except Exception as e:
            logger.error(f"Error generating PDF: {str(e)}")
            raise BadRequest(f"Error generating PDF: {str(e)}")

    def export_as_docx(self, form_data: Dict[str, Any], format_params: Dict[str, Any]) -> bytes:
        """Export form as fillable DOCX with custom formatting"""
        try:
            self._validate_form_data(form_data)

            doc = Document()
            
            # Set page size and margins
            section = doc.sections[0]
            if format_params['page_size'] == 'A4':
                section.page_width = Inches(8.27)
                section.page_height = Inches(11.69)
            elif format_params['page_size'] == 'LEGAL':
                section.page_width = Inches(8.5)
                section.page_height = Inches(14)
            else:  # LETTER
                section.page_width = Inches(8.5)
                section.page_height = Inches(11)

            # Set margins
            section.left_margin = Inches(format_params['margin_left'])
            section.right_margin = Inches(format_params['margin_right'])
            section.top_margin = Inches(format_params['margin_top'])
            section.bottom_margin = Inches(format_params['margin_bottom'])

            # Add logo if provided
            if format_params.get('logo_path'):
                self._add_logo_docx(doc, format_params['logo_path'])

            # Set default font size and line spacing
            style = doc.styles['Normal']
            style.font.size = Pt(format_params['font_size'])
            style.paragraph_format.line_spacing = format_params['line_spacing']

            # Title
            title = doc.add_heading(form_data['title'], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER

            # Description
            if form_data.get('description'):
                desc = doc.add_paragraph(form_data['description'])
                desc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            doc.add_paragraph()

            # Form Information
            doc.add_heading('Form Information:', level=1)
            doc.add_paragraph(f"Environment: {form_data['created_by']['environment']['name']}")
            doc.add_paragraph(f"Created by: {form_data['created_by']['fullname']}")
            doc.add_paragraph()

            # Response Information
            doc.add_heading('Response Information:', level=1)
            doc.add_paragraph("Name: _________________________________")
            doc.add_paragraph("Date: _________________________________")
            doc.add_paragraph()

            # Questions
            for i, question in enumerate(form_data['questions'], 1):
                # Question text
                p = doc.add_paragraph()
                p.add_run(f"{i}. {question['text']}").bold = True
                
                # Handle different question types
                if question['type'] == 'text':
                    doc.add_paragraph("_________________________________")
                
                elif question['type'] in ['checkbox', 'multiple_choices']:
                    if question.get('possible_answers'):
                        for answer in question['possible_answers']:
                            p = doc.add_paragraph()
                            p.add_run(f"□ {answer['value']}")
                            p.style = 'List Bullet'
                    else:
                        doc.add_paragraph("□ Yes    □ No")

                # Add remarks if present
                if question.get('remarks'):
                    p = doc.add_paragraph()
                    p.add_run("Remarks: ").bold = True
                    p.add_run(question['remarks'])

                doc.add_paragraph()

            # Add signatures
            self._add_signatures_docx(doc, format_params.get('signatures', []))

            # Save to buffer
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            return buffer.getvalue()

        except Exception as e:
            logger.error(f"Error generating DOCX: {str(e)}")
            raise BadRequest(f"Error generating DOCX: {str(e)}")

    @staticmethod
    def get_supported_formats() -> List[str]:
        """Get list of supported export formats"""
        return ['PDF', 'DOCX']

    def validate_format(self, format: str) -> None:
        """Validate export format"""
        if format.upper() not in self.supported_formats:
            raise ValueError(
                f"Unsupported format: {format}. Supported formats: {', '.join(self.supported_formats)}"
            )
            
    def _validate_signature_spacing(self, spacing: dict) -> tuple[bool, Optional[str]]:
        """Validate signature spacing parameters"""
        try:
            valid_ranges = {
                'before_section': (5, 50),
                'between_signatures': (5, 30),
                'after_date': (2, 20),
                'after_section': (5, 50)
            }
            
            for key, (min_val, max_val) in valid_ranges.items():
                if key in spacing:
                    value = float(spacing[key])
                    if not min_val <= value <= max_val:
                        return False, f"Signature spacing {key} must be between {min_val} and {max_val}"
                        
            return True, None
        except ValueError:
            return False, "Invalid signature spacing values"