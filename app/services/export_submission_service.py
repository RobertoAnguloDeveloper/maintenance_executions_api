# app/services/export_submission_service.py

import io
import warnings
import itertools
from typing import Dict, List, Optional, Tuple, Union, Any
from datetime import datetime
import os
import logging
from io import BytesIO
from sqlalchemy.orm import joinedload
from reportlab.lib import colors
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle as ReportLabParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_RIGHT, TA_JUSTIFY
from reportlab.lib.units import inch, mm
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph as ReportLabParagraph, Spacer,
    Table as ReportLabTable, TableStyle as ReportLabTableStyle,
    Image as ReportLabImage, Flowable, KeepTogether
)
from PIL import Image as PILImage # type: ignore
import json
import re
from collections import defaultdict

from werkzeug.datastructures import FileStorage # type: ignore

# Assuming your models are correctly imported from your app structure
# Models like FormQuestion, Question, QuestionType are no longer directly used for Q&A content
from app.models.form_submission import FormSubmission
from app.models.attachment import Attachment
from app.models.answer_submitted import AnswerSubmitted # Primary source for Q&A
from app.models.form import Form


# Imports for DOCX generation
from docx import Document # type: ignore
from docx.shared import Inches, Pt, RGBColor # type: ignore
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING # type: ignore
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT # type: ignore
from docx.oxml import OxmlElement # type: ignore
from docx.oxml.ns import qn, nsdecls # type: ignore
from docx.oxml import parse_xml # type: ignore


logger = logging.getLogger(__name__)
warnings.filterwarnings("ignore", category=UserWarning, module="docx.styles.styles", message="style lookup by style_id is deprecated")

# --- Default Style Configuration (for PDF and adapted for DOCX) ---
DEFAULT_STYLE_CONFIG: Dict[str, Any] = {
    # Page Layout
    "page_margin_top": 0.75, "page_margin_bottom": 0.75, "page_margin_left": 0.75, "page_margin_right": 0.75,
    "default_font_family": "Helvetica", "default_font_family_docx": "Calibri",
    "default_font_color": colors.black, "default_font_color_docx": "000000", "default_font_size_docx": 11,
    # Title
    "title_font_family": "Helvetica-Bold", "title_font_family_docx": "Calibri", "title_font_size": 18, "title_font_size_docx": 22,
    "title_leading": 22, "title_font_color": colors.black, "title_font_color_docx": "000000",
    "title_alignment": TA_CENTER, "title_alignment_docx": "center",
    "title_space_after": 0.25, "title_space_after_docx": 12,
    # Description
    "description_font_family": "Helvetica", "description_font_family_docx": "Calibri", "description_font_size": 10, "description_font_size_docx": 10,
    "description_leading": 12, "description_font_color": colors.darkslategray, "description_font_color_docx": "2F4F4F",
    "description_alignment": TA_LEFT, "description_alignment_docx": "left",
    "description_space_after": 0.15, "description_space_after_docx": 6,
    # Submission Info
    "info_font_family": "Helvetica", "info_font_family_docx": "Calibri", "info_font_size": 10, "info_font_size_docx": 9,
    "info_leading": 12, "info_font_color": colors.darkslategray, "info_font_color_docx": "2F4F4F",
    "info_label_font_family": "Helvetica-Bold", "info_label_font_family_docx": "Calibri",
    "info_alignment": TA_LEFT, "info_alignment_docx": "left",
    "info_space_after": 0.2, "info_space_after_docx": 12,
    # Question
    "question_font_family": "Helvetica-Bold", "question_font_family_docx": "Calibri", "question_font_size": 11, "question_font_size_docx": 12,
    "question_font_color": colors.black, "question_font_color_docx": "000000",
    "question_left_indent": 0, "question_left_indent_docx": 0, # Inches for both, service layer adapts if needed
    "question_space_before": 0.15, "question_space_before_docx": 8, # Inches for PDF, Points for DOCX
    "question_space_after": 4/72.0, "question_space_after_docx": 3, # Inches (effectively points) for PDF, Points for DOCX
    "question_leading": 14,
    # Answer
    "answer_font_family": "Helvetica", "answer_font_family_docx": "Calibri", "answer_font_size": 10, "answer_font_size_docx": 10,
    "answer_font_color": colors.darkslategray, "answer_font_color_docx": "2F4F4F",
    "answer_left_indent": 15/72.0, "answer_left_indent_docx": 0.25, # Inches (effectively points) for PDF, Inches for DOCX
    "answer_space_before": 2/72.0, "answer_space_before_docx": 2, # Inches (effectively points) for PDF, Points for DOCX
    "answer_space_after": 0.15, "answer_space_after_docx": 6, # Inches for PDF, Points for DOCX
    "answer_leading": 12, "qa_layout": "answer_below", "answer_same_line_max_length": 70,
    # Table Header
    "table_header_font_family": "Helvetica-Bold", "table_header_font_family_docx": "Calibri", "table_header_font_size": 9, "table_header_font_size_docx": 10,
    "table_header_font_color": colors.black, "table_header_font_color_docx": "000000",
    "table_header_bg_color": colors.lightgrey, "table_header_bg_color_docx": "D3D3D3",
    "table_header_leading": 11, "table_header_alignment": "CENTER", "table_header_alignment_docx": "center",
    "table_cell_padding_left": 3, "table_cell_padding_right": 3, "table_cell_padding_top": 3, "table_cell_padding_bottom": 3, # Points for ReportLab
    "table_space_after": 0.15, "table_space_after_docx": 12, # Inches for PDF, Points for DOCX
    # Table Cell
    "table_cell_font_family": "Helvetica", "table_cell_font_family_docx": "Calibri", "table_cell_font_size": 8, "table_cell_font_size_docx": 9,
    "table_cell_font_color": colors.black, "table_cell_font_color_docx": "000000",
    "table_cell_leading": 10, "table_cell_alignment": "LEFT", "table_cell_alignment_docx": "left",
    "table_grid_color": colors.grey, "table_grid_color_docx": "BEBEBE", "table_grid_thickness": 0.5, # Points
    # Signatures
    "signature_label_font_family": "Helvetica-Bold", "signature_label_font_family_docx": "Calibri", 
    "signature_label_font_size": 12, "signature_label_font_size_docx": 11,
    "signature_label_font_color": colors.black, "signature_label_font_color_docx": "000000",
    "signature_text_font_family": "Helvetica", "signature_text_font_family_docx": "Calibri", 
    "signature_text_font_size": 9, "signature_text_font_size_docx": 9,
    "signature_text_font_color": colors.black, "signature_text_font_color_docx": "000000",
    "signature_image_width": 2.0, "signature_image_width_docx": 2.0, # Inches
    "signature_image_height": 0.8, "signature_image_height_docx": 0.8, # Inches
    "signature_section_space_before": 0.3, "signature_section_space_before_docx": 18, # Inches for PDF, Points for DOCX
    "signature_space_between_vertical": 0.2, "signature_space_between_vertical_docx_pt": 10, # Inches for PDF, Points for DOCX
    "signature_position_alignment": "center", "signature_position_alignment_docx": "center",  # NEW: left, center, right
}

# --- Global Helper Functions ---
def _get_color_rl(color_input: Any, default_color: colors.Color = colors.black) -> colors.Color:
    if isinstance(color_input, colors.Color): return color_input
    if isinstance(color_input, str):
        try:
            if color_input.startswith("#"): return colors.HexColor(color_input)
            color_name_lower = color_input.lower()
            if hasattr(colors, color_name_lower): return getattr(colors, color_name_lower)
            if len(color_input) == 6 and all(c in "0123456789abcdefABCDEF" for c in color_input): return colors.HexColor(f"#{color_input}")
            logger.warning(f"Unknown ReportLab color string '{color_input}'. Using default.")
            return default_color
        except (ValueError, AttributeError): logger.warning(f"Invalid ReportLab color string '{color_input}'. Using default."); return default_color
    elif isinstance(color_input, (list, tuple)) and len(color_input) == 3:
        try: return colors.Color(float(color_input[0]), float(color_input[1]), float(color_input[2]))
        except ValueError: logger.warning(f"Invalid RGB list/tuple for color: {color_input}. Using default."); return default_color
    logger.warning(f"Invalid ReportLab color type '{type(color_input)}'. Using default.")
    return default_color

def _color_to_hex_string_rl(color_obj: Any) -> str: # Returns hex string without '#'
    if isinstance(color_obj, str):
        try:
            temp_color = _get_color_rl(color_obj) # Ensure it's a ReportLab Color object
            if hasattr(temp_color, 'hexval') and callable(getattr(temp_color, 'hexval')):
                  # hexval() often includes '#', so strip it.
                return temp_color.hexval().lstrip('#')
            # Fallback for non-HexColor ReportLab Color objects
            r_float, g_float, b_float = temp_color.red, temp_color.green, temp_color.blue
            return f"{int(r_float*255):02x}{int(g_float*255):02x}{int(b_float*255):02x}"
        except: return "000000"
    elif hasattr(color_obj, 'red') and hasattr(color_obj, 'green') and hasattr(color_obj, 'blue'): # Assuming it's a ReportLab color
        r_float = max(0.0, min(1.0, color_obj.red)); g_float = max(0.0, min(1.0, color_obj.green)); b_float = max(0.0, min(1.0, color_obj.blue))
        return f"{int(r_float*255):02x}{int(g_float*255):02x}{int(b_float*255):02x}"
    return "000000"

def _parse_numeric_value(value: Any, default_value: float = 0.0) -> float:
    if value is None: return default_value
    try: return float(value)
    except (ValueError, TypeError): logger.warning(f"Invalid numeric value '{value}'. Using default {default_value}."); return default_value

def _get_alignment_code_rl(align_input: Optional[Union[str, int]], default_align: int = TA_LEFT) -> int:
    code_map = {"LEFT": TA_LEFT, "CENTER": TA_CENTER, "RIGHT": TA_RIGHT, "JUSTIFY": TA_JUSTIFY}
    if isinstance(align_input, int) and align_input in code_map.values(): return align_input
    if isinstance(align_input, str):
        effective_align_str = align_input.upper()
        if effective_align_str in code_map: return code_map[effective_align_str]
    return default_align

def _get_docx_alignment(align_input: Optional[Union[str, int]], default_align: WD_ALIGN_PARAGRAPH = WD_ALIGN_PARAGRAPH.LEFT) -> WD_ALIGN_PARAGRAPH:
    align_map_str = {"LEFT": WD_ALIGN_PARAGRAPH.LEFT, "CENTER": WD_ALIGN_PARAGRAPH.CENTER, "RIGHT": WD_ALIGN_PARAGRAPH.RIGHT, "JUSTIFY": WD_ALIGN_PARAGRAPH.JUSTIFY}
    align_map_int = {TA_LEFT: WD_ALIGN_PARAGRAPH.LEFT, 0: WD_ALIGN_PARAGRAPH.LEFT, TA_CENTER: WD_ALIGN_PARAGRAPH.CENTER, 1: WD_ALIGN_PARAGRAPH.CENTER, TA_RIGHT: WD_ALIGN_PARAGRAPH.RIGHT, 2: WD_ALIGN_PARAGRAPH.RIGHT, TA_JUSTIFY: WD_ALIGN_PARAGRAPH.JUSTIFY, 4: WD_ALIGN_PARAGRAPH.JUSTIFY}
    if isinstance(align_input, str): return align_map_str.get(align_input.upper(), default_align)
    if isinstance(align_input, int): return align_map_int.get(align_input, default_align)
    return default_align

def _get_docx_color(color_input: Any, default_hex: str = "000000") -> Optional[RGBColor]:
    hex_val = default_hex.lstrip('#').upper() # Ensure default is 6-char hex
    if isinstance(color_input, str):
        color_str_cleaned = color_input.lower().strip()
        if color_str_cleaned.startswith("#"): hex_val = color_str_cleaned[1:].upper()
        else:
            color_name_map = {"black": "000000", "white": "FFFFFF", "red": "FF0000", "green": "00FF00", "blue": "0000FF", "yellow": "FFFF00", "cyan": "00FFFF", "magenta": "FF00FF", "silver": "C0C0C0", "gray": "808080", "grey": "808080", "maroon": "800000", "olive": "808000", "purple": "800080", "teal": "008080", "navy": "000080", "darkblue": "00008B", "darkgrey": "A9A9A9", "darkgray": "A9A9A9", "lightgrey": "D3D3D3", "lightgray": "D3D3D3", "darkslategray":"2F4F4F"}
            if color_str_cleaned in color_name_map: hex_val = color_name_map[color_str_cleaned]
            elif len(color_str_cleaned) == 6 and all(c in "0123456789ABCDEF" for c in color_str_cleaned.upper()): hex_val = color_str_cleaned.upper()
            else: logger.warning(f"DOCX: Unknown color string '{color_input}'. Using default hex '{default_hex}'."); hex_val = default_hex.lstrip('#').upper()
    elif isinstance(color_input, RGBColor): return color_input # Already a docx RGBColor
    elif hasattr(color_input, 'red') and hasattr(color_input, 'green') and hasattr(color_input, 'blue'): # ReportLab color
        try:
            r = int(max(0.0, min(1.0, color_input.red)) * 255)
            g = int(max(0.0, min(1.0, color_input.green)) * 255)
            b = int(max(0.0, min(1.0, color_input.blue)) * 255)
            return RGBColor(r, g, b)
        except Exception as e: logger.warning(f"Could not convert ReportLab color to DOCX RGBColor: {e}. Using default."); hex_val = default_hex.lstrip('#').upper()
    
    if len(hex_val) == 3: hex_val = "".join([c*2 for c in hex_val]) # Expand shorthand hex
    if len(hex_val) != 6: logger.warning(f"DOCX: Invalid hex color value '{hex_val}' derived from '{color_input}'. Using default."); hex_val = default_hex.lstrip('#').upper()
    
    try: return RGBColor(int(hex_val[0:2], 16), int(hex_val[2:4], 16), int(hex_val[4:6], 16))
    except ValueError:
        logger.warning(f"DOCX: Error converting hex '{hex_val}' to RGBColor. Using default.")
        default_hex_clean = default_hex.lstrip('#').upper()
        def_r, def_g, def_b = int(default_hex_clean[0:2],16), int(default_hex_clean[2:4],16), int(default_hex_clean[4:6],16)
        return RGBColor(def_r, def_g, def_b)

def _set_cell_background_docx(cell, color_string_input: str):
    docx_rgb_object = _get_docx_color(color_string_input) 
    if docx_rgb_object is None: 
        logger.warning(f"DOCX: _get_docx_color returned None for background color input '{color_string_input}'. Skipping background.")
        return
    hex_fill_value = str(docx_rgb_object) 
    shading_elm_str = f'<w:shd {nsdecls("w")} w:fill="{hex_fill_value}" w:val="clear" />'
    try:
        shading_elm = parse_xml(shading_elm_str)
        tcPr = cell._tc.get_or_add_tcPr()
        for existing_shd in tcPr.xpath('w:shd'): tcPr.remove(existing_shd)
        tcPr.append(shading_elm)
    except Exception as e: logger.error(f"Error setting DOCX cell background with color {hex_fill_value}: {e}", exc_info=True)


class ExportSubmissionService:
    @staticmethod
    def _process_header_image(
        image_file: Union[FileStorage, str],
        upload_path_base: Optional[str] = None,
        size_percent: Optional[float] = None,
        width_px: Optional[float] = None,
        height_px: Optional[float] = None
    ) -> Optional[Tuple[BytesIO, float, float]]:
        try:
            img_data = None
            if isinstance(image_file, FileStorage):
                img_data = image_file.read()
                image_file.seek(0) 
            elif isinstance(image_file, str):
                full_path = os.path.join(upload_path_base, image_file) if upload_path_base and not os.path.isabs(image_file) else image_file
                if os.path.exists(full_path):
                    with open(full_path, 'rb') as f: img_data = f.read()
                else: logger.error(f"Header image file not found at path: {full_path}"); return None
            else: logger.error(f"Invalid header_image type: {type(image_file)}"); return None

            if not img_data: return None
            img = PILImage.open(io.BytesIO(img_data))
            if img.format == 'GIF' and getattr(img, 'is_animated', False) and getattr(img, 'n_frames', 1) > 1:
                img.seek(0) 

            orig_width_px, orig_height_px = img.size
            new_width_px, new_height_px = float(orig_width_px), float(orig_height_px)

            if width_px is not None and height_px is not None: new_width_px, new_height_px = float(width_px), float(height_px)
            elif width_px is not None:
                new_width_px = float(width_px)
                new_height_px = new_width_px * (orig_height_px / orig_width_px) if orig_width_px > 0 else 0
            elif height_px is not None:
                new_height_px = float(height_px)
                new_width_px = new_height_px * (orig_width_px / orig_height_px) if orig_height_px > 0 else 0
            elif size_percent is not None:
                scale_factor = float(size_percent) / 100.0
                new_width_px *= scale_factor; new_height_px *= scale_factor
            
            if new_width_px <=0 or new_height_px <=0: new_width_px, new_height_px = float(orig_width_px), float(orig_height_px)

            if img.mode in ('P', 'LA') or (img.mode == 'RGBA' and img.info.get('transparency') is not None): img = img.convert('RGBA')
            elif img.mode != 'RGB': img = img.convert('RGB')
            img_resized = img.resize((int(new_width_px), int(new_height_px)), PILImage.LANCZOS)
            
            img_format_upper = img.format.upper() if img.format else 'PNG'
            if img_format_upper in ['SVG', 'WEBP', 'GIF']: img_format_upper = 'PNG'
            result_io = io.BytesIO()
            if img_format_upper == 'JPEG' and img_resized.mode == 'RGBA': img_resized = img_resized.convert('RGB')
            
            img_resized.save(result_io, format=img_format_upper)
            result_io.seek(0)
            return result_io, new_width_px, new_height_px
        except Exception as e: logger.error(f"Error processing header image: {str(e)}", exc_info=True); return None

    @staticmethod
    def _add_header_image_to_reportlab_story(
        story: List[Flowable], header_image_input: Union[FileStorage, str], upload_path_base: str,
        opacity: float, size_percent: Optional[float], width_px: Optional[float], height_px: Optional[float],
        alignment_str: str, page_margin_left: float, page_margin_right: float, doc_width: float
    ):
        processed_image_data = ExportSubmissionService._process_header_image(header_image_input, upload_path_base, size_percent, width_px, height_px)
        if processed_image_data:
            image_io, img_w_px, img_h_px = processed_image_data
            img_w_pt = img_w_px * (72.0 / 96.0); img_h_pt = img_h_px * (72.0 / 96.0)
            available_width_for_image = doc_width
            if img_w_pt > available_width_for_image:
                scale_ratio = available_width_for_image / img_w_pt
                img_w_pt *= scale_ratio; img_h_pt *= scale_ratio
            if opacity < 1.0: logger.warning("ReportLab Image opacity < 1.0 not directly supported.")
            img_obj = ReportLabImage(image_io, width=img_w_pt, height=img_h_pt)
            align_code = _get_alignment_code_rl(alignment_str.upper(), TA_CENTER)
            table_style_align_map = {TA_LEFT: 'LEFT', TA_CENTER: 'CENTER', TA_RIGHT: 'RIGHT'}
            img_align_for_table_style = table_style_align_map.get(align_code, 'CENTER')
            img_table = ReportLabTable([[img_obj]], colWidths=[doc_width])
            img_table.setStyle(ReportLabTableStyle([('ALIGN', (0,0), (0,0), img_align_for_table_style), ('VALIGN', (0,0), (0,0), 'MIDDLE'), ('LEFTPADDING', (0,0), (0,0), 0), ('RIGHTPADDING', (0,0), (0,0), 0), ('TOPPADDING', (0,0), (0,0), 0), ('BOTTOMPADDING', (0,0), (0,0), 0)]))
            story.append(img_table); story.append(Spacer(1, 0.1 * inch))

    @staticmethod
    def _add_signatures_to_reportlab_story(
        story: List[Flowable], attachments: List[Attachment], upload_path: str,
        scale_factor: float, alignment_str: str, position_alignment: str,  # NEW parameter
        styles: Any, config: Dict[str, Any]
    ):
        sig_attachments = [att for att in attachments if att.is_signature and not att.is_deleted]
        if not sig_attachments: return

        if 'CustomSignatureLabel' not in styles:
            styles.add(ReportLabParagraphStyle(name='CustomSignatureLabel', 
                fontName=str(config.get("signature_label_font_family", "Helvetica-Bold")), 
                fontSize=_parse_numeric_value(config.get("signature_label_font_size", 12)), 
                textColor=_get_color_rl(config.get("signature_label_font_color", colors.black)), 
                spaceBefore=_parse_numeric_value(config.get("signature_section_space_before", 0.3)) * inch, 
                spaceAfter=4))
        if 'CustomSignatureText' not in styles:
            styles.add(ReportLabParagraphStyle(name='CustomSignatureText', 
                fontName=str(config.get("signature_text_font_family", "Helvetica")), 
                fontSize=_parse_numeric_value(config.get("signature_text_font_size", 9)), 
                textColor=_get_color_rl(config.get("signature_text_font_color", colors.black)), 
                leading=_parse_numeric_value(config.get("signature_text_font_size", 9)) * 1.2))
        
        story.append(ReportLabParagraph("Signatures:", styles['CustomSignatureLabel']))
        story.append(Spacer(1, 0.05 * inch))

        sig_img_w_conf = _parse_numeric_value(config.get("signature_image_width", 2.0)) * inch * scale_factor
        sig_img_h_conf = _parse_numeric_value(config.get("signature_image_height", 0.8)) * inch * scale_factor
        sig_space_between_vertical = _parse_numeric_value(config.get("signature_space_between_vertical", 0.2)) * inch
        page_content_width = styles.PAGE_WIDTH - styles.LEFT_MARGIN - styles.RIGHT_MARGIN

        # Map position alignment to ReportLab table alignment
        position_align_map = {
            "left": "LEFT",
            "center": "CENTER", 
            "right": "RIGHT"
        }
        table_position_align = position_align_map.get(position_alignment.lower(), "CENTER")

        if alignment_str.lower() == "horizontal" and len(sig_attachments) > 1:
            effective_sig_width = sig_img_w_conf + (0.2 * inch) 
            max_sigs_per_row = int(page_content_width / effective_sig_width) if effective_sig_width > 0 else 1
            max_sigs_per_row = max(1, min(max_sigs_per_row, len(sig_attachments), 3)) 
            sig_rows_data: List[List[List[Flowable]]] = []
            current_sig_row_items: List[List[Flowable]] = []
            
            for idx, att in enumerate(sig_attachments):
                sig_block_elements: List[Flowable] = []
                full_file_path = os.path.join(upload_path, att.file_path)
                sig_author = att.signature_author or "N/A"
                sig_position = att.signature_position or "N/A"
                
                if os.path.exists(full_file_path):
                    try: 
                        sig_block_elements.append(ReportLabImage(full_file_path, width=sig_img_w_conf, height=sig_img_h_conf))
                    except Exception as e_img: 
                        sig_block_elements.append(ReportLabParagraph(f"<i>[Image Error: {e_img}]</i>", styles['CustomSignatureText']))
                else: 
                    sig_block_elements.append(ReportLabParagraph(f"<i>[Image Missing: {os.path.basename(att.file_path)}]</i>", styles['CustomSignatureText']))
                
                sig_block_elements.extend([
                    Spacer(1, 2), 
                    ReportLabParagraph("___________________________", styles['CustomSignatureText']), 
                    ReportLabParagraph(f"<b>Signed by:</b> {sig_author}", styles['CustomSignatureText']), 
                    ReportLabParagraph(f"<b>Position:</b> {sig_position}", styles['CustomSignatureText'])
                ])
                current_sig_row_items.append(sig_block_elements)
                
                if len(current_sig_row_items) == max_sigs_per_row or idx == len(sig_attachments) - 1:
                    sig_rows_data.append(current_sig_row_items)
                    current_sig_row_items = []
            
            for sig_row_group in sig_rows_data:
                if not sig_row_group: continue
                
                # Calculate column width based on position alignment
                if position_alignment.lower() == "left":
                    col_widths = [effective_sig_width] * len(sig_row_group)
                elif position_alignment.lower() == "right":
                    remaining_width = page_content_width - (effective_sig_width * len(sig_row_group))
                    col_widths = [remaining_width] + [effective_sig_width] * len(sig_row_group)
                    sig_row_group = [[]] + sig_row_group  # Add empty first column
                else:  # center
                    col_width_sig = page_content_width / len(sig_row_group) if len(sig_row_group) > 0 else page_content_width
                    col_widths = [col_width_sig] * len(sig_row_group)
                
                sig_table = ReportLabTable([sig_row_group], colWidths=col_widths)
                sig_table.setStyle(ReportLabTableStyle([
                    ('ALIGN', (0,0), (-1,-1), table_position_align),
                    ('VALIGN', (0,0), (-1,-1), 'TOP'), 
                    ('LEFTPADDING', (0,0), (-1,-1), 2), 
                    ('RIGHTPADDING', (0,0), (-1,-1), 2), 
                    ('BOTTOMPADDING', (0,0), (-1,-1), 5)
                ]))
                story.append(sig_table)
                story.append(Spacer(1, 0.1 * inch))
        else:  # Vertical alignment
            # For vertical alignment, create a container table for position alignment
            for att in sig_attachments:
                full_file_path = os.path.join(upload_path, att.file_path)
                sig_author = att.signature_author or "N/A"
                sig_position = att.signature_position or "N/A"
                
                sig_elements = []
                
                if os.path.exists(full_file_path):
                    try: 
                        sig_elements.append(ReportLabImage(full_file_path, width=sig_img_w_conf, height=sig_img_h_conf))
                    except Exception as e_img: 
                        sig_elements.append(ReportLabParagraph(f"<i>[Sig Image Error: {e_img}]</i>", styles['CustomSignatureText']))
                else: 
                    sig_elements.append(ReportLabParagraph(f"<i>[Sig Image Missing: {os.path.basename(att.file_path)}]</i>", styles['CustomSignatureText']))
                
                sig_elements.extend([
                    Spacer(1, 2), 
                    ReportLabParagraph("___________________________", styles['CustomSignatureText']), 
                    ReportLabParagraph(f"<b>Signed by:</b> {sig_author}", styles['CustomSignatureText']), 
                    ReportLabParagraph(f"<b>Position:</b> {sig_position}", styles['CustomSignatureText'])
                ])
                
                # Wrap in a table for alignment control
                sig_container = ReportLabTable([[sig_elements]], colWidths=[page_content_width])
                sig_container.setStyle(ReportLabTableStyle([
                    ('ALIGN', (0,0), (0,0), table_position_align),
                    ('VALIGN', (0,0), (0,0), 'TOP'),
                    ('LEFTPADDING', (0,0), (0,0), 0),
                    ('RIGHTPADDING', (0,0), (0,0), 0),
                    ('TOPPADDING', (0,0), (0,0), 0),
                    ('BOTTOMPADDING', (0,0), (0,0), 0)
                ]))
                story.append(sig_container)
                story.append(Spacer(1, sig_space_between_vertical))

    @staticmethod
    def _add_header_image_to_docx_header(
        doc: Document, header_image_input: Union[FileStorage, str], upload_path_base: str,
        size_percent: Optional[float], width_px: Optional[float], height_px: Optional[float],
        alignment_str: str
    ):
        processed_image_data = ExportSubmissionService._process_header_image(header_image_input, upload_path_base, size_percent, width_px, height_px)
        if processed_image_data:
            image_io, img_w_px, _ = processed_image_data
            img_width_inches = img_w_px / 96.0 
            header = doc.sections[0].header
            for para in header.paragraphs: para._p.getparent().remove(para._p)
            p_header_img = header.add_paragraph()
            run_header_img = p_header_img.add_run()
            try: run_header_img.add_picture(image_io, width=Inches(img_width_inches))
            except Exception as e: logger.error(f"DOCX: Error adding header picture: {e}"); run_header_img.add_text("[Header Image Error]")
            p_header_img.alignment = _get_docx_alignment(alignment_str, WD_ALIGN_PARAGRAPH.CENTER)

    @staticmethod
    def _add_signatures_to_docx(
        doc: Document, attachments: List[Attachment], upload_path: str,
        size_percent: float, alignment_str: str, position_alignment: str,  # NEW parameter
        config: Dict[str, Any]
    ):
        sig_attachments_info = [
            {
                'path': os.path.join(upload_path, att.file_path),
                'author': att.signature_author or "N/A",
                'position': att.signature_position or "N/A",
                'exists': os.path.exists(os.path.join(upload_path, att.file_path))
            } 
            for att in attachments if att.is_signature and not att.is_deleted
        ]
        if not sig_attachments_info: return
        
        sig_heading_p = doc.add_paragraph()
        sig_heading_run = sig_heading_p.add_run("Signatures")
        sig_heading_run.bold = True
        sig_heading_run.font.name = str(config.get("signature_label_font_family_docx", "Calibri"))
        sig_heading_run.font.size = Pt(_parse_numeric_value(config.get("signature_label_font_size_docx", 11)))
        sig_label_font_color_val = _get_docx_color(config.get("signature_label_font_color_docx", "000000"))
        if sig_label_font_color_val: 
            sig_heading_run.font.color.rgb = sig_label_font_color_val
        sig_heading_p.paragraph_format.space_before = Pt(_parse_numeric_value(config.get("signature_section_space_before_docx", 18)))
        sig_heading_p.paragraph_format.space_after = Pt(6)

        sig_img_width_default_inches = _parse_numeric_value(config.get("signature_image_width_docx", 2.0))
        sig_img_width_final_inches = sig_img_width_default_inches * (size_percent / 100.0)
        sig_text_font_fam = str(config.get("signature_text_font_family_docx", "Calibri"))
        sig_text_font_sz = Pt(_parse_numeric_value(config.get("signature_text_font_size_docx", 9)))
        sig_text_font_color = _get_docx_color(config.get("signature_text_font_color_docx", "000000"))
        sig_space_after_pt = Pt(_parse_numeric_value(config.get("signature_space_between_vertical_docx_pt", 10)))

        # Map position alignment to DOCX alignment
        docx_position_align = _get_docx_alignment(position_alignment, WD_ALIGN_PARAGRAPH.CENTER)

        if alignment_str.lower() == "horizontal" and len(sig_attachments_info) > 1:
            page_width_inches = (doc.sections[0].page_width - doc.sections[0].left_margin - doc.sections[0].right_margin) / Inches(1)
            effective_sig_block_width = sig_img_width_final_inches + 0.5
            max_cols = int(page_width_inches / effective_sig_block_width) if effective_sig_block_width > 0 else 1
            max_cols = max(1, min(max_cols, len(sig_attachments_info), 3))
            num_rows = (len(sig_attachments_info) + max_cols - 1) // max_cols
            
            sig_table = doc.add_table(rows=num_rows, cols=max_cols)
            
            # Set table alignment based on position_alignment
            if position_alignment.lower() == "left":
                sig_table.alignment = WD_TABLE_ALIGNMENT.LEFT
            elif position_alignment.lower() == "right":
                sig_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
            else:  # center
                sig_table.alignment = WD_TABLE_ALIGNMENT.CENTER
            
            sig_table.style = 'TableGrid'
            
            # Remove borders
            for row_obj in sig_table.rows:
                for cell_obj in row_obj.cells:
                    tcPr = cell_obj._tc.get_or_add_tcPr()
                    tcBorders = OxmlElement('w:tcBorders')
                    for border_name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                        border_el = OxmlElement(f'w:{border_name}')
                        border_el.set(qn('w:val'), 'nil')
                        tcBorders.append(border_el)
                    tcPr.append(tcBorders)

            for idx, sig_info in enumerate(sig_attachments_info):
                cell = sig_table.cell(idx // max_cols, idx % max_cols)
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
                if cell.paragraphs and cell.paragraphs[0].text == "":
                    p_elem = cell.paragraphs[0]._element
                    p_elem.getparent().remove(p_elem)

                # Add signature elements with consistent alignment
                p_img = cell.add_paragraph()
                p_img.alignment = docx_position_align
                if sig_info["exists"]:
                    try: 
                        p_img.add_run().add_picture(sig_info["path"], width=Inches(sig_img_width_final_inches))
                    except Exception as e: 
                        p_img.add_run(f"[Img Err: {e}]")
                else: 
                    p_img.add_run(f"[Img Miss: {os.path.basename(sig_info['path'])}]")
                
                p_line = cell.add_paragraph("___________________________")
                p_line.alignment = docx_position_align
                for run in p_line.runs: 
                    run.font.name = sig_text_font_fam
                    run.font.size = sig_text_font_sz
                    if sig_text_font_color: 
                        run.font.color.rgb = sig_text_font_color
                
                p_author = cell.add_paragraph()
                p_author.alignment = docx_position_align
                author_run = p_author.add_run("Signed by: ")
                author_run.bold = True
                p_author.add_run(sig_info['author'])
                for run in p_author.runs: 
                    run.font.name = sig_text_font_fam
                    run.font.size = sig_text_font_sz
                    if sig_text_font_color: 
                        run.font.color.rgb = sig_text_font_color
                
                p_position = cell.add_paragraph()
                p_position.alignment = docx_position_align
                pos_run = p_position.add_run("Position: ")
                pos_run.bold = True
                p_position.add_run(sig_info['position'])
                for run in p_position.runs: 
                    run.font.name = sig_text_font_fam
                    run.font.size = sig_text_font_sz
                    if sig_text_font_color: 
                        run.font.color.rgb = sig_text_font_color
                p_position.paragraph_format.space_after = Pt(0)
            
            doc.add_paragraph().paragraph_format.space_after = sig_space_after_pt
        
        else:  # Vertical alignment
            for sig_info in sig_attachments_info:
                p_img = doc.add_paragraph()
                p_img.alignment = docx_position_align
                if sig_info["exists"]:
                    try: 
                        p_img.add_run().add_picture(sig_info["path"], width=Inches(sig_img_width_final_inches))
                    except Exception as e: 
                        p_img.add_run(f"[Img Err: {e}]")
                else: 
                    p_img.add_run(f"[Img Miss: {os.path.basename(sig_info['path'])}]")
                
                p_line = doc.add_paragraph("___________________________")
                p_line.alignment = docx_position_align
                for run in p_line.runs: 
                    run.font.name = sig_text_font_fam
                    run.font.size = sig_text_font_sz
                    if sig_text_font_color: 
                        run.font.color.rgb = sig_text_font_color

                p_author = doc.add_paragraph()
                p_author.alignment = docx_position_align
                author_run = p_author.add_run("Signed by: ")
                author_run.bold = True
                p_author.add_run(sig_info['author'])
                for run in p_author.runs: 
                    run.font.name = sig_text_font_fam
                    run.font.size = sig_text_font_sz
                    if sig_text_font_color: 
                        run.font.color.rgb = sig_text_font_color

                p_position = doc.add_paragraph()
                p_position.alignment = docx_position_align
                pos_run = p_position.add_run("Position: ")
                pos_run.bold = True
                p_position.add_run(sig_info['position'])
                for run in p_position.runs: 
                    run.font.name = sig_text_font_fam
                    run.font.size = sig_text_font_sz
                    if sig_text_font_color: 
                        run.font.color.rgb = sig_text_font_color
                p_position.paragraph_format.space_after = sig_space_after_pt

    # ============================================================================
    # == MAIN EXPORT METHODS 
    # ============================================================================
    @staticmethod
    def export_structured_submission_to_pdf(
        submission_id: int,
        upload_path: str,
        include_signatures: bool = True,
        header_image: Optional[Any] = None,
        header_opacity: float = 1.0,
        header_size: Optional[float] = None,
        header_width: Optional[float] = None,
        header_height: Optional[float] = None,
        header_alignment: str = "center",
        signatures_size: float = 100,
        signatures_alignment: str = "vertical",
        signatures_position_alignment: str = "left",  # NEW parameter
        pdf_style_options: Optional[Dict[str, Any]] = None
    ) -> Tuple[Optional[BytesIO], Optional[str]]:
        
        final_config = DEFAULT_STYLE_CONFIG.copy()
        user_overrides = {} 
        if pdf_style_options:
            for key, value in pdf_style_options.items():
                if key in final_config:
                    final_config[key] = value 
                    user_overrides[key] = value 
                else: logger.warning(f"PDF Export: Unknown style option '{key}' provided.")

        buffer = BytesIO()
        try:
            submission = FormSubmission.query.options(
                joinedload(FormSubmission.form), 
                joinedload(FormSubmission.answers_submitted),
                joinedload(FormSubmission.attachments)
            ).filter_by(id=submission_id, is_deleted=False).first()

            if not submission: return None, "Submission not found"
            form = submission.form
            if not form: return None, "Form not found for submission"

            page_margin_top_val = _parse_numeric_value(final_config.get("page_margin_top"), 0.75) * inch
            page_margin_bottom_val = _parse_numeric_value(final_config.get("page_margin_bottom"), 0.75) * inch
            page_margin_left_val = _parse_numeric_value(final_config.get("page_margin_left"), 0.75) * inch
            page_margin_right_val = _parse_numeric_value(final_config.get("page_margin_right"), 0.75) * inch

            doc_pdf = SimpleDocTemplate(buffer, pagesize=letter,
                                        leftMargin=page_margin_left_val, rightMargin=page_margin_right_val,
                                        topMargin=page_margin_top_val, bottomMargin=page_margin_bottom_val)
            story: List[Flowable] = []
            styles_pdf = getSampleStyleSheet()

            pdf_spacing_keys_in_inches_default = ["title_space_after", "description_space_after", "info_space_after", "question_left_indent", "question_space_before", "question_space_after", "answer_left_indent", "answer_space_before", "answer_space_after", "table_space_after"]
            def get_style_attr_val_points(key_name: str, default_config_value_if_not_overridden: Any) -> float:
                if key_name in user_overrides: return _parse_numeric_value(user_overrides[key_name])
                else: return _parse_numeric_value(default_config_value_if_not_overridden) * inch

            styles_pdf.add(ReportLabParagraphStyle(name='CustomTitle', fontName=str(final_config["title_font_family"]), fontSize=_parse_numeric_value(final_config["title_font_size"]), leading=_parse_numeric_value(final_config.get("title_leading", _parse_numeric_value(final_config["title_font_size"]) * 1.2)), textColor=_get_color_rl(final_config["title_font_color"]), alignment=_get_alignment_code_rl(final_config["title_alignment"]), spaceAfter=get_style_attr_val_points("title_space_after", DEFAULT_STYLE_CONFIG["title_space_after"])))
            styles_pdf.add(ReportLabParagraphStyle(name='CustomDescription', fontName=str(final_config.get("description_font_family","Helvetica")), fontSize=_parse_numeric_value(final_config.get("description_font_size",10)), leading=_parse_numeric_value(final_config.get("description_leading",12)), textColor=_get_color_rl(final_config.get("description_font_color",colors.darkslategray)), alignment=_get_alignment_code_rl(final_config.get("description_alignment",TA_LEFT)), spaceAfter=get_style_attr_val_points("description_space_after", DEFAULT_STYLE_CONFIG["description_space_after"])))
            styles_pdf.add(ReportLabParagraphStyle(name='CustomSubmissionInfo', fontName=str(final_config["info_font_family"]), fontSize=_parse_numeric_value(final_config["info_font_size"]), leading=_parse_numeric_value(final_config.get("info_leading", _parse_numeric_value(final_config["info_font_size"]) * 1.2)), textColor=_get_color_rl(final_config["info_font_color"]), alignment=_get_alignment_code_rl(final_config.get("info_alignment", TA_LEFT)), spaceAfter=get_style_attr_val_points("info_space_after", DEFAULT_STYLE_CONFIG["info_space_after"])))
            styles_pdf.add(ReportLabParagraphStyle(name='CustomQuestion', fontName=str(final_config["question_font_family"]), fontSize=_parse_numeric_value(final_config["question_font_size"]), leading=_parse_numeric_value(final_config["question_leading"]), textColor=_get_color_rl(final_config["question_font_color"]), leftIndent=get_style_attr_val_points("question_left_indent", DEFAULT_STYLE_CONFIG["question_left_indent"]), spaceBefore=get_style_attr_val_points("question_space_before", DEFAULT_STYLE_CONFIG["question_space_before"]), spaceAfter=get_style_attr_val_points("question_space_after", DEFAULT_STYLE_CONFIG["question_space_after"])))
            styles_pdf.add(ReportLabParagraphStyle(name='CustomAnswer', fontName=str(final_config["answer_font_family"]), fontSize=_parse_numeric_value(final_config["answer_font_size"]), leading=_parse_numeric_value(final_config["answer_leading"]), textColor=_get_color_rl(final_config["answer_font_color"]), leftIndent=get_style_attr_val_points("answer_left_indent", DEFAULT_STYLE_CONFIG["answer_left_indent"]), spaceBefore=get_style_attr_val_points("answer_space_before", DEFAULT_STYLE_CONFIG["answer_space_before"]), spaceAfter=get_style_attr_val_points("answer_space_after", DEFAULT_STYLE_CONFIG["answer_space_after"])))
            styles_pdf.add(ReportLabParagraphStyle(name='CustomQACombined', fontName=str(final_config["question_font_family"]), fontSize=_parse_numeric_value(final_config["question_font_size"]), leading=_parse_numeric_value(final_config["question_leading"]), textColor=_get_color_rl(final_config["question_font_color"]), leftIndent=get_style_attr_val_points("question_left_indent", DEFAULT_STYLE_CONFIG["question_left_indent"]), spaceBefore=get_style_attr_val_points("question_space_before", DEFAULT_STYLE_CONFIG["question_space_before"]), spaceAfter=get_style_attr_val_points("answer_space_after", DEFAULT_STYLE_CONFIG["answer_space_after"])))
            styles_pdf.add(ReportLabParagraphStyle(name='CustomTableHeader', fontName=str(final_config["table_header_font_family"]), fontSize=_parse_numeric_value(final_config["table_header_font_size"]), textColor=_get_color_rl(final_config["table_header_font_color"]), alignment=_get_alignment_code_rl(final_config["table_header_alignment"]), leading=_parse_numeric_value(final_config.get("table_header_leading",_parse_numeric_value(final_config["table_header_font_size"])*1.2))))
            styles_pdf.add(ReportLabParagraphStyle(name='CustomTableCell', fontName=str(final_config["table_cell_font_family"]), fontSize=_parse_numeric_value(final_config["table_cell_font_size"]), textColor=_get_color_rl(final_config["table_cell_font_color"]), alignment=_get_alignment_code_rl(final_config["table_cell_alignment"]), leading=_parse_numeric_value(final_config.get("table_cell_leading",_parse_numeric_value(final_config["table_cell_font_size"])*1.2))))
            styles_pdf.add(ReportLabParagraphStyle(name='CustomTableError', parent=styles_pdf['CustomTableCell'], textColor=colors.red))
            styles_pdf.PAGE_WIDTH = letter[0]; styles_pdf.LEFT_MARGIN = page_margin_left_val; styles_pdf.RIGHT_MARGIN = page_margin_right_val
            
            if header_image: ExportSubmissionService._add_header_image_to_reportlab_story(story, header_image, upload_path, header_opacity, header_size, header_width, header_height, header_alignment, doc_pdf.leftMargin, doc_pdf.rightMargin, doc_pdf.width)
            if form.title: story.append(ReportLabParagraph(form.title, styles_pdf['CustomTitle']))
            if form.description: story.append(ReportLabParagraph(form.description.replace('\n','<br/>\n') if form.description else "", styles_pdf['CustomDescription']))
            info_text = f"<b>Submitted by:</b> {submission.submitted_by or 'N/A'}<br/><b>Date:</b> {submission.submitted_at.strftime('%Y-%m-%d %H:%M:%S') if submission.submitted_at else 'N/A'}"
            story.append(ReportLabParagraph(info_text, styles_pdf['CustomSubmissionInfo']))

            # --- Data processing from answers_submitted ---
            all_renderable_items = []
            answers_by_question_text_and_order = defaultdict(lambda: {"text": "", "type": "", "order": float('inf'), "answers_list": []})
            cell_based_tables_data = defaultdict(lambda: {'name':'', 'headers':{}, 'cells':{}, 'row_indices':set(), 'col_indices':set(), 'header_row_present':False, 'order':float('inf')})

            # Initial pass to populate structures
            for ans in submission.answers_submitted:
                if not (ans.question and ans.question_type): continue
                
                # Replace tilde in question text here
                q_text_key = str(ans.question).replace("~", " ").strip()
                q_type_lower = str(ans.question_type).lower().strip()
                q_order = ans.question_order if ans.question_order is not None else float('inf')

                if q_type_lower == 'table':
                    cell_based_tables_data[q_text_key]['name'] = q_text_key
                    cell_based_tables_data[q_text_key]['order'] = min(cell_based_tables_data[q_text_key]['order'], q_order)
                    cell_content_str = str(ans.cell_content if ans.cell_content is not None else ans.answer or "").strip()
                    try:
                        if ans.row is None or ans.column is None: logger.warning(f"PDF: Table cell for '{q_text_key}' has None for row or column. Skipping."); continue
                        row_idx, col_idx = int(ans.row), int(ans.column)
                        if row_idx == 0: 
                            cell_based_tables_data[q_text_key]['headers'][col_idx] = cell_content_str
                            cell_based_tables_data[q_text_key]['col_indices'].add(col_idx)
                            cell_based_tables_data[q_text_key]['header_row_present'] = True
                        elif row_idx > 0: 
                            cell_based_tables_data[q_text_key]['cells'][(row_idx,col_idx)] = cell_content_str
                            cell_based_tables_data[q_text_key]['row_indices'].add(row_idx)
                            cell_based_tables_data[q_text_key]['col_indices'].add(col_idx)
                    except (ValueError,TypeError) as e_cell: logger.warning(f"PDF: Invalid table cell index for table '{q_text_key}' (row='{ans.row}', col='{ans.column}', error: {e_cell}). Skipping cell."); continue
                elif q_type_lower != 'signature': # Regular Q&A
                    unique_q_key = (q_order, q_text_key)
                    if not answers_by_question_text_and_order[unique_q_key]["text"]: # Initialize if first time
                         answers_by_question_text_and_order[unique_q_key]["text"] = q_text_key
                         answers_by_question_text_and_order[unique_q_key]["type"] = q_type_lower
                         answers_by_question_text_and_order[unique_q_key]["order"] = q_order
                    answers_by_question_text_and_order[unique_q_key]["answers_list"].append(ans.answer)
            
            # Prepare Q&A items
            for _key, q_data in answers_by_question_text_and_order.items():
                ans_val_display = "<i>No answer provided</i>"
                q_type_lower = q_data['type']
                raw_answers = q_data['answers_list']

                if q_type_lower in ['dropdown','select','multiselect','checkbox','multiple_choices','single_choice']:
                    combined_options=[]
                    for ans_content in raw_answers:
                        if ans_content is not None and str(ans_content).strip() != "":
                            try:
                                parsed_json_options=json.loads(ans_content)
                                if isinstance(parsed_json_options,list): combined_options.extend([str(item).strip() for item in parsed_json_options if str(item).strip()])
                                elif str(parsed_json_options).strip(): combined_options.append(str(parsed_json_options).strip())
                            except (json.JSONDecodeError,TypeError):
                                if str(ans_content).strip(): combined_options.append(str(ans_content).strip())
                    unique_options=list(dict.fromkeys(opt for opt in combined_options if opt))
                    ans_val_display=", ".join(unique_options) if unique_options else "<i>No selection</i>"
                else: 
                    ans_val_display = ", ".join(str(a).strip() for a in raw_answers if a is not None and str(a).strip()) or "<i>No answer provided</i>"

                all_renderable_items.append({
                    "type": "qa", "order": q_data['order'], 
                    "question_text": q_data['text'], "answer_display": ans_val_display
                })

            # Prepare Table items
            for table_name, table_render_data in cell_based_tables_data.items():
                all_renderable_items.append({
                    "type": "table", "order": table_render_data['order'],
                    "data": table_render_data, "name": table_name 
                })
            
            all_renderable_items.sort(key=lambda x: (x['order'] if x['order'] is not None else float('inf'), x.get('name', x.get('question_text', ''))))

            qa_layout = str(final_config.get("qa_layout","answer_below"))
            answer_same_line_max_len = int(_parse_numeric_value(final_config.get("answer_same_line_max_length"),70))
            answer_font_color_html = _color_to_hex_string_rl(final_config["answer_font_color"])

            for item in all_renderable_items:
                if item['type'] == "qa":
                    q_text_p_escaped = item['question_text'].replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('\n','<br/>\n')
                    ans_val_p_escaped = item['answer_display'].replace('&','&amp;').replace('<','&lt;').replace('>','&gt;').replace('\n','<br/>\n')
                    
                    if qa_layout == "answer_same_line" and len(item['answer_display']) <= answer_same_line_max_len and \
                       ('<br/>' not in ans_val_p_escaped and '\n' not in item['answer_display']): 
                        combined_text_html = f"<b>{q_text_p_escaped}:</b> <font color='#{answer_font_color_html}'>{ans_val_p_escaped}</font>"
                        story.append(ReportLabParagraph(combined_text_html, styles_pdf['CustomQACombined']))
                    else:
                        q_paragraph = ReportLabParagraph(q_text_p_escaped, styles_pdf['CustomQuestion'])
                        a_paragraph = ReportLabParagraph(ans_val_p_escaped, styles_pdf['CustomAnswer'])
                        story.append(KeepTogether([q_paragraph, a_paragraph]))

                elif item['type'] == "table":
                    table_data_render = item['data']
                    q_text = table_data_render['name'] 
                    story.append(ReportLabParagraph(q_text.replace('\n','<br/>\n'), styles_pdf['CustomQuestion']))

                    if table_data_render and (table_data_render['header_row_present'] or table_data_render['cells']):
                        all_cols = sorted(list(table_data_render['col_indices']))
                        if not all_cols: 
                            story.append(ReportLabParagraph("<i>Table has no columns defined or no data.</i>",styles_pdf['CustomTableError']))
                            story.append(Spacer(1,get_style_attr_val_points("table_space_after", DEFAULT_STYLE_CONFIG["table_space_after"])))
                            continue
                        
                        table_rows_styled_content=[]
                        if table_data_render['header_row_present']: 
                            header_row=[ReportLabParagraph(str(table_data_render['headers'].get(c,'')).replace('\n','<br/>\n'),styles_pdf['CustomTableHeader']) for c in all_cols]
                            table_rows_styled_content.append(header_row)
                        
                        data_row_indices = sorted(list(r for r in table_data_render['row_indices'] if r > 0))
                        for r_idx in data_row_indices: 
                            data_row=[ReportLabParagraph(str(table_data_render['cells'].get((r_idx,c),'')).replace('\n','<br/>\n'),styles_pdf['CustomTableCell']) for c in all_cols]
                            table_rows_styled_content.append(data_row)
                        
                        if not table_rows_styled_content: story.append(ReportLabParagraph("<i>No data rows for this table.</i>",styles_pdf['CustomAnswer']))
                        else:
                            rl_table_obj=ReportLabTable(table_rows_styled_content,repeatRows=(1 if table_data_render['header_row_present'] else 0))
                            table_style_cmds=[('GRID',(0,0),(-1,-1),_parse_numeric_value(final_config["table_grid_thickness"]),_get_color_rl(final_config["table_grid_color"])),('VALIGN',(0,0),(-1,-1),'MIDDLE'),('LEFTPADDING',(0,0),(-1,-1),_parse_numeric_value(final_config["table_cell_padding_left"])),('RIGHTPADDING',(0,0),(-1,-1),_parse_numeric_value(final_config["table_cell_padding_right"])),('TOPPADDING',(0,0),(-1,-1),_parse_numeric_value(final_config["table_cell_padding_top"])),('BOTTOMPADDING',(0,0),(-1,-1),_parse_numeric_value(final_config["table_cell_padding_bottom"]))]
                            if table_data_render['header_row_present']: table_style_cmds.append(('BACKGROUND',(0,0),(-1,0),_get_color_rl(final_config["table_header_bg_color"])))
                            rl_table_obj.setStyle(ReportLabTableStyle(table_style_cmds)); story.append(rl_table_obj)
                    else: story.append(ReportLabParagraph("<i>No data submitted for this table.</i>",styles_pdf['CustomAnswer']))
                    story.append(Spacer(1, get_style_attr_val_points("table_space_after", DEFAULT_STYLE_CONFIG["table_space_after"])))
            
            if include_signatures: ExportSubmissionService._add_signatures_to_reportlab_story(
                story, submission.attachments, upload_path, 
                signatures_size/100.0, signatures_alignment, 
                signatures_position_alignment,  # NEW parameter
                styles_pdf, final_config
            )
            
            doc_pdf.build(story); buffer.seek(0); return buffer,None
        except NameError as ne:
            logger.error(f"Error exporting structured submission to PDF: {submission_id} - {str(ne)}", exc_info=True); error_message=f"An error occurred during PDF generation: {str(ne)}. Ensure 'KeepTogether' is imported from 'reportlab.platypus'."
            if "KeepTogether" in str(ne): logger.error("Critical Import Missing: from reportlab.platypus import KeepTogether"); error_message="Critical Import Missing: from reportlab.platypus import KeepTogether. Please add this to your script."
            return None,error_message
        except Exception as e: logger.error(f"Error exporting structured submission to PDF: {submission_id} - {str(e)}", exc_info=True); return None,f"An error occurred during PDF generation: {str(e)}"

    @staticmethod
    def export_submission_to_pdf(*args, **kwargs): # Wrapper for backward compatibility
        return ExportSubmissionService.export_structured_submission_to_pdf(*args, **kwargs)

    @staticmethod
    def export_submission_to_docx(
        submission_id: int, 
        upload_path: str, 
        include_signatures: bool = True,
        style_options: Optional[Dict[str, Any]] = None, 
        header_image_file: Optional[Any] = None,
        header_size_percent: Optional[float] = None, 
        header_width_px: Optional[float] = None,
        header_height_px: Optional[float] = None, 
        header_alignment_str: str = "center",
        signatures_size_percent: float = 100, 
        signatures_alignment_str: str = "vertical",
        signatures_position_alignment_str: str = "left"
    ) -> Tuple[Optional[BytesIO], Optional[str]]:
        final_config_docx = DEFAULT_STYLE_CONFIG.copy()
        if style_options:
            for key, value in style_options.items(): final_config_docx[key] = value 

        try:
            submission = FormSubmission.query.options(
                joinedload(FormSubmission.form), 
                joinedload(FormSubmission.answers_submitted), 
                joinedload(FormSubmission.attachments)
            ).filter_by(id=submission_id, is_deleted=False).first()

            if not submission: return None, "Submission not found"
            form = submission.form
            if not form: return None, "Form not found for submission"

            doc = Document()
            style = doc.styles['Normal']; font = style.font 
            font.name = str(final_config_docx.get("default_font_family_docx", "Calibri"))
            font.size = Pt(_parse_numeric_value(final_config_docx.get("default_font_size_docx", 11)))
            default_font_color_val = _get_docx_color(final_config_docx.get("default_font_color_docx", "000000"))
            if default_font_color_val: font.color.rgb = default_font_color_val # type: ignore
            
            for section in doc.sections:
                section.top_margin = Inches(_parse_numeric_value(final_config_docx.get("page_margin_top"), 0.75))
                section.bottom_margin = Inches(_parse_numeric_value(final_config_docx.get("page_margin_bottom"), 0.75))
                section.left_margin = Inches(_parse_numeric_value(final_config_docx.get("page_margin_left"), 0.75))
                section.right_margin = Inches(_parse_numeric_value(final_config_docx.get("page_margin_right"), 0.75))

            if header_image_file: ExportSubmissionService._add_header_image_to_docx_header(doc, header_image_file, upload_path, header_size_percent, header_width_px, header_height_px, header_alignment_str)

            if form.title:
                title_p = doc.add_paragraph(); title_run = title_p.add_run(form.title)
                title_run.font.name = str(final_config_docx.get("title_font_family_docx", font.name)); title_run.font.size = Pt(_parse_numeric_value(final_config_docx.get("title_font_size_docx", 22))); title_run.bold = True
                title_font_color_val = _get_docx_color(final_config_docx.get("title_font_color_docx", "000000"))
                if title_font_color_val: title_run.font.color.rgb = title_font_color_val # type: ignore
                title_p.alignment = _get_docx_alignment(final_config_docx.get("title_alignment_docx", "center"))
                title_p.paragraph_format.space_after = Pt(_parse_numeric_value(final_config_docx.get("title_space_after_docx", 12)))
            
            if form.description:
                desc_p = doc.add_paragraph()
                desc_lines = (form.description or "").split('\n')
                for i, line in enumerate(desc_lines):
                    desc_p.add_run(line)
                    if i < len(desc_lines) - 1: desc_p.add_run().add_break()
                
                for run in desc_p.runs:
                    run.font.name = str(final_config_docx.get("description_font_family_docx", font.name)); run.font.size = Pt(_parse_numeric_value(final_config_docx.get("description_font_size_docx", 10)))
                    desc_font_color_val = _get_docx_color(final_config_docx.get("description_font_color_docx", "2F4F4F"))
                    if desc_font_color_val: run.font.color.rgb = desc_font_color_val # type: ignore
                desc_p.alignment = _get_docx_alignment(final_config_docx.get("description_alignment_docx", "left")); desc_p.paragraph_format.space_after = Pt(_parse_numeric_value(final_config_docx.get("description_space_after_docx", 6)))

            info_p = doc.add_paragraph()
            info_label_font_name = str(final_config_docx.get("info_label_font_family_docx", font.name)) 
            info_text_font_name = str(final_config_docx.get("info_font_family_docx", font.name))
            info_font_sz_val = Pt(_parse_numeric_value(final_config_docx.get("info_font_size_docx", 9)))
            info_font_color_val = _get_docx_color(final_config_docx.get("info_font_color_docx", "2F4F4F"))

            run_submitted_by_label = info_p.add_run("Submitted by: "); run_submitted_by_label.font.name = info_label_font_name; run_submitted_by_label.font.size = info_font_sz_val; run_submitted_by_label.bold = True
            if info_font_color_val: run_submitted_by_label.font.color.rgb = info_font_color_val # type: ignore
            run_submitted_by_val = info_p.add_run(f"{submission.submitted_by or 'N/A'}\n"); run_submitted_by_val.font.name = info_text_font_name; run_submitted_by_val.font.size = info_font_sz_val
            if info_font_color_val: run_submitted_by_val.font.color.rgb = info_font_color_val # type: ignore
            run_date_label = info_p.add_run("Date: "); run_date_label.font.name = info_label_font_name; run_date_label.font.size = info_font_sz_val; run_date_label.bold = True
            if info_font_color_val: run_date_label.font.color.rgb = info_font_color_val # type: ignore
            run_date_val = info_p.add_run(submission.submitted_at.strftime('%Y-%m-%d %H:%M:%S') if submission.submitted_at else 'N/A'); run_date_val.font.name = info_text_font_name; run_date_val.font.size = info_font_sz_val
            if info_font_color_val: run_date_val.font.color.rgb = info_font_color_val # type: ignore
            info_p.alignment = _get_docx_alignment(final_config_docx.get("info_alignment_docx", "left")); info_p.paragraph_format.space_after = Pt(_parse_numeric_value(final_config_docx.get("info_space_after_docx", 12)))

            # --- Data processing from answers_submitted for DOCX ---
            all_renderable_items_docx = []
            answers_by_question_text_and_order_docx = defaultdict(lambda: {"text": "", "type": "", "order": float('inf'), "answers_list": []})
            cell_based_tables_data_docx = defaultdict(lambda: {'name':'', 'headers':{}, 'cells':{}, 'row_indices':set(), 'col_indices':set(), 'header_row_present':False, 'order':float('inf')})

            for ans_docx in submission.answers_submitted:
                if not (ans_docx.question and ans_docx.question_type): continue
                
                # Replace tilde in question text here
                q_text_key_docx = str(ans_docx.question).replace("~", " ").strip()
                q_type_lower_docx = str(ans_docx.question_type).lower().strip()
                q_order_docx = ans_docx.question_order if ans_docx.question_order is not None else float('inf')

                if q_type_lower_docx == 'table':
                    cell_based_tables_data_docx[q_text_key_docx]['name'] = q_text_key_docx
                    cell_based_tables_data_docx[q_text_key_docx]['order'] = min(cell_based_tables_data_docx[q_text_key_docx]['order'], q_order_docx)
                    cell_content_str_docx = str(ans_docx.cell_content if ans_docx.cell_content is not None else ans_docx.answer or "").strip()
                    try:
                        if ans_docx.row is None or ans_docx.column is None: logger.warning(f"DOCX: Table cell for '{q_text_key_docx}' has None for row or column. Skipping."); continue
                        row_idx_docx, col_idx_docx = int(ans_docx.row), int(ans_docx.column)
                        if row_idx_docx == 0: 
                            cell_based_tables_data_docx[q_text_key_docx]['headers'][col_idx_docx] = cell_content_str_docx
                            cell_based_tables_data_docx[q_text_key_docx]['col_indices'].add(col_idx_docx)
                            cell_based_tables_data_docx[q_text_key_docx]['header_row_present'] = True
                        elif row_idx_docx > 0: 
                            cell_based_tables_data_docx[q_text_key_docx]['cells'][(row_idx_docx,col_idx_docx)] = cell_content_str_docx
                            cell_based_tables_data_docx[q_text_key_docx]['row_indices'].add(row_idx_docx)
                            cell_based_tables_data_docx[q_text_key_docx]['col_indices'].add(col_idx_docx)
                    except (ValueError, TypeError) as e_cell_docx: logger.warning(f"DOCX: Invalid table cell index for table '{q_text_key_docx}' (row='{ans_docx.row}', col='{ans_docx.column}', error: {e_cell_docx}). Skipping cell."); continue
                elif q_type_lower_docx != 'signature':
                    unique_q_key_docx = (q_order_docx, q_text_key_docx)
                    if not answers_by_question_text_and_order_docx[unique_q_key_docx]["text"]:
                        answers_by_question_text_and_order_docx[unique_q_key_docx]["text"] = q_text_key_docx
                        answers_by_question_text_and_order_docx[unique_q_key_docx]["type"] = q_type_lower_docx
                        answers_by_question_text_and_order_docx[unique_q_key_docx]["order"] = q_order_docx
                    answers_by_question_text_and_order_docx[unique_q_key_docx]["answers_list"].append(ans_docx.answer)

            for _key_docx, q_data_docx in answers_by_question_text_and_order_docx.items():
                ans_text_val_docx = "No answer provided"
                q_type_val_docx = q_data_docx['type']
                raw_answers_docx = q_data_docx['answers_list']
                
                if q_type_val_docx in ['dropdown', 'select', 'multiselect', 'checkbox', 'multiple_choices', 'single_choice']:
                    combined_options_docx = []
                    for ans_content_docx in raw_answers_docx:
                        if ans_content_docx is not None and str(ans_content_docx).strip() != "":
                            try:
                                parsed_json_options_docx = json.loads(ans_content_docx)
                                if isinstance(parsed_json_options_docx, list): combined_options_docx.extend([str(item).strip() for item in parsed_json_options_docx if str(item).strip()])
                                elif str(parsed_json_options_docx).strip(): combined_options_docx.append(str(parsed_json_options_docx).strip())
                            except (json.JSONDecodeError, TypeError):
                                if str(ans_content_docx).strip(): combined_options_docx.append(str(ans_content_docx).strip())
                    unique_options_docx = list(dict.fromkeys(opt for opt in combined_options_docx if opt))
                    ans_text_val_docx = ", ".join(unique_options_docx) if unique_options_docx else "No selection"
                else:
                    ans_text_val_docx = ", ".join(str(a).strip() for a in raw_answers_docx if a is not None and str(a).strip()) or "No answer provided"
                
                all_renderable_items_docx.append({
                    "type": "qa", "order": q_data_docx['order'],
                    "question_text": q_data_docx['text'], "answer_display": ans_text_val_docx
                })

            for table_name_docx, table_render_data_docx in cell_based_tables_data_docx.items():
                all_renderable_items_docx.append({
                    "type": "table", "order": table_render_data_docx['order'],
                    "data": table_render_data_docx, "name": table_name_docx
                })
            
            all_renderable_items_docx.sort(key=lambda x: (x['order'] if x['order'] is not None else float('inf'), x.get('name', x.get('question_text', ''))))
            
            qa_layout_pref_docx = str(final_config_docx.get("qa_layout", "answer_below"))
            ans_same_line_max_len_docx = int(_parse_numeric_value(final_config_docx.get("answer_same_line_max_length"), 70))

            for item_docx in all_renderable_items_docx:
                if item_docx['type'] == "qa":
                    q_text_val_docx = item_docx['question_text']
                    ans_text_val_docx = item_docx['answer_display']

                    if qa_layout_pref_docx == "answer_same_line" and len(ans_text_val_docx) <= ans_same_line_max_len_docx and '\n' not in ans_text_val_docx:
                        p_qa = doc.add_paragraph(); q_run = p_qa.add_run(f"{q_text_val_docx}: ")
                        q_run.font.name = str(final_config_docx.get("question_font_family_docx", font.name)); q_run.font.size = Pt(_parse_numeric_value(final_config_docx.get("question_font_size_docx", 12))); q_run.bold = True
                        q_font_color_val = _get_docx_color(final_config_docx.get("question_font_color_docx", "000000"))
                        if q_font_color_val: q_run.font.color.rgb = q_font_color_val # type: ignore
                        a_run = p_qa.add_run(ans_text_val_docx)
                        a_run.font.name = str(final_config_docx.get("answer_font_family_docx", font.name)); a_run.font.size = Pt(_parse_numeric_value(final_config_docx.get("answer_font_size_docx", 10)))
                        ans_font_color_val = _get_docx_color(final_config_docx.get("answer_font_color_docx", "2F4F4F"))
                        if ans_font_color_val: a_run.font.color.rgb = ans_font_color_val # type: ignore
                        p_qa.paragraph_format.space_before = Pt(_parse_numeric_value(final_config_docx.get("question_space_before_docx", 8)))
                        p_qa.paragraph_format.space_after = Pt(_parse_numeric_value(final_config_docx.get("answer_space_after_docx", 6)))
                    else:
                        q_p = doc.add_paragraph()
                        q_lines = q_text_val_docx.split('\n')
                        for i_qline, q_line in enumerate(q_lines): q_p.add_run(q_line); 
                        if i_qline < len(q_lines) -1: q_p.add_run().add_break()
                        for run in q_p.runs: 
                            run.font.name = str(final_config_docx.get("question_font_family_docx", font.name)); run.font.size = Pt(_parse_numeric_value(final_config_docx.get("question_font_size_docx", 12))); run.bold = True
                            q_font_color_val = _get_docx_color(final_config_docx.get("question_font_color_docx", "000000"))
                            if q_font_color_val: run.font.color.rgb = q_font_color_val # type: ignore
                        q_p.paragraph_format.space_before = Pt(_parse_numeric_value(final_config_docx.get("question_space_before_docx", 8))); q_p.paragraph_format.space_after = Pt(_parse_numeric_value(final_config_docx.get("question_space_after_docx", 3)))
                        
                        a_p = doc.add_paragraph()
                        a_lines = ans_text_val_docx.split('\n')
                        for i_aline, a_line in enumerate(a_lines): a_p.add_run(a_line)
                        if i_aline < len(a_lines) -1: a_p.add_run().add_break()
                        for run in a_p.runs: 
                            run.font.name = str(final_config_docx.get("answer_font_family_docx", font.name)); run.font.size = Pt(_parse_numeric_value(final_config_docx.get("answer_font_size_docx", 10)))
                            ans_font_color_val = _get_docx_color(final_config_docx.get("answer_font_color_docx", "2F4F4F"))
                            if ans_font_color_val: run.font.color.rgb = ans_font_color_val # type: ignore
                        a_p.paragraph_format.left_indent = Inches(_parse_numeric_value(final_config_docx.get("answer_left_indent_docx", 0.25)))
                        a_p.paragraph_format.space_after = Pt(_parse_numeric_value(final_config_docx.get("answer_space_after_docx", 6)))
                
                elif item_docx['type'] == "table":
                    table_data_render_docx = item_docx['data']
                    q_text_val_docx = table_data_render_docx['name']
                    tbl_title_p_docx = doc.add_paragraph(); tbl_title_run_docx = tbl_title_p_docx.add_run(q_text_val_docx.split('\n')[0])
                    tbl_title_run_docx.font.name = str(final_config_docx.get("question_font_family_docx", font.name)); tbl_title_run_docx.font.size = Pt(_parse_numeric_value(final_config_docx.get("question_font_size_docx", 12))); tbl_title_run_docx.bold = True
                    q_font_color_val = _get_docx_color(final_config_docx.get("question_font_color_docx", "000000"))
                    if q_font_color_val: tbl_title_run_docx.font.color.rgb = q_font_color_val # type: ignore
                    tbl_title_p_docx.paragraph_format.space_before = Pt(_parse_numeric_value(final_config_docx.get("question_space_before_docx", 8))); tbl_title_p_docx.paragraph_format.space_after = Pt(_parse_numeric_value(final_config_docx.get("question_space_after_docx", 3)))

                    if table_data_render_docx and (table_data_render_docx['header_row_present'] or table_data_render_docx['cells']):
                        all_cols_docx = sorted(list(table_data_render_docx['col_indices']))
                        if not all_cols_docx: doc.add_paragraph("Table has no columns defined or no data.").paragraph_format.space_after = Pt(_parse_numeric_value(final_config_docx.get("table_space_after_docx", 12))); continue
                        
                        data_row_indices_docx = sorted(list(r for r in table_data_render_docx['row_indices'] if r > 0))
                        num_data_rows_docx = len(data_row_indices_docx)
                        total_docx_rows = (1 if table_data_render_docx['header_row_present'] else 0) + num_data_rows_docx
                        
                        if total_docx_rows == 0: doc.add_paragraph("No data for this table.").paragraph_format.space_after = Pt(_parse_numeric_value(final_config_docx.get("table_space_after_docx", 12))); continue
                        
                        docx_table_obj = doc.add_table(rows=total_docx_rows, cols=len(all_cols_docx)); docx_table_obj.style = str(final_config_docx.get('table_style_docx', 'TableGrid'))
                        table_align_str_val = final_config_docx.get("table_alignment_docx", "left") 
                        if table_align_str_val == "center": docx_table_obj.alignment = WD_TABLE_ALIGNMENT.CENTER
                        elif table_align_str_val == "right": docx_table_obj.alignment = WD_TABLE_ALIGNMENT.RIGHT
                        else: docx_table_obj.alignment = WD_TABLE_ALIGNMENT.LEFT
                        
                        if table_data_render_docx['header_row_present']:
                            header_cells_docx = docx_table_obj.rows[0].cells
                            for c_idx, actual_col_idx in enumerate(all_cols_docx):
                                cell_p = header_cells_docx[c_idx].paragraphs[0]; cell_p.text = str(table_data_render_docx['headers'].get(actual_col_idx, '')); cell_p.alignment = _get_docx_alignment(final_config_docx.get("table_header_alignment_docx", "center"))
                                for run in cell_p.runs: 
                                    run.font.name = str(final_config_docx.get("table_header_font_family_docx", font.name)); run.font.size = Pt(_parse_numeric_value(final_config_docx.get("table_header_font_size_docx", 10))); run.bold = True
                                    th_font_color_val = _get_docx_color(final_config_docx.get("table_header_font_color_docx", "000000"))
                                    if th_font_color_val: run.font.color.rgb = th_font_color_val #type: ignore
                                _set_cell_background_docx(header_cells_docx[c_idx], str(final_config_docx.get("table_header_bg_color_docx", "D3D3D3")))
                        
                        current_docx_data_row_render_idx = 1 if table_data_render_docx['header_row_present'] else 0
                        for r_form_idx in data_row_indices_docx: 
                            if current_docx_data_row_render_idx >= total_docx_rows: break 
                            row_cells_docx = docx_table_obj.rows[current_docx_data_row_render_idx].cells
                            for c_idx, actual_col_idx in enumerate(all_cols_docx):
                                cell_p_data = row_cells_docx[c_idx].paragraphs[0]; cell_p_data.text = str(table_data_render_docx['cells'].get((r_form_idx, actual_col_idx), '')); cell_p_data.alignment = _get_docx_alignment(final_config_docx.get("table_cell_alignment_docx", "left"))
                                for run in cell_p_data.runs: 
                                    run.font.name = str(final_config_docx.get("table_cell_font_family_docx", font.name)); run.font.size = Pt(_parse_numeric_value(final_config_docx.get("table_cell_font_size_docx", 9)))
                                    tc_font_color_val = _get_docx_color(final_config_docx.get("table_cell_font_color_docx", "000000"))
                                    if tc_font_color_val: run.font.color.rgb = tc_font_color_val #type: ignore
                            current_docx_data_row_render_idx += 1
                        doc.add_paragraph().paragraph_format.space_after = Pt(_parse_numeric_value(final_config_docx.get("table_space_after_docx", 12)))
                    else: 
                        doc.add_paragraph("No data submitted for this table.").paragraph_format.space_after = Pt(_parse_numeric_value(final_config_docx.get("table_space_after_docx", 12)))

            if include_signatures:
                ExportSubmissionService._add_signatures_to_docx(
                doc, submission.attachments, upload_path, 
                signatures_size_percent, signatures_alignment_str, 
                signatures_position_alignment_str,  # NEW parameter
                final_config_docx
            )

            buffer_docx = BytesIO()
            doc.save(buffer_docx)
            buffer_docx.seek(0)
            return buffer_docx, None
        except Exception as e:
            logger.error(f"Error exporting submission {submission_id} to DOCX: {str(e)}", exc_info=True)
            return None, f"An error occurred during DOCX generation: {str(e)}"