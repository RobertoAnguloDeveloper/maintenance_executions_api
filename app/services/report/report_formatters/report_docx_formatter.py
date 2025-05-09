# app/services/report/report_formatters/report_docx_formatter.py
from typing import Dict, Any, List, Optional
from io import BytesIO
import logging
from datetime import datetime
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import qn

from ..report_formatter import ReportFormatter

logger = logging.getLogger(__name__)

class ReportDocxFormatter(ReportFormatter):
    """Formatter for DOCX reports"""
    
    def generate(self) -> BytesIO:
        """
        Generate a DOCX report
        
        Returns:
            BytesIO buffer with the DOCX report
        """
        document = docx.Document()
        buffer = BytesIO()
        
        # Set document properties
        document.core_properties.title = self.report_title
        document.core_properties.created = datetime.now()
        
        # Set up default styles
        style = document.styles['Normal']
        style.font.name = 'Calibri'
        style.font.size = Pt(11)
        
        # Add report title
        title = document.add_heading(self.report_title, level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Add generation timestamp
        timestamp = document.add_paragraph(f"Generated: {self.generation_timestamp}")
        timestamp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        timestamp.style.font.italic = True
        
        document.add_paragraph()  # Spacing
        
        # Process each report type
        first_section = True
        for report_type, result in self.processed_data.items():
            # Skip internal data
            if report_type.startswith('_'):
                continue
                
            # Handle errors
            if result.get('error'):
                p = document.add_paragraph()
                p.add_run(f"Error: {report_type}:").italic = True
                p = document.add_paragraph()
                p.add_run(str(result['error'])).italic = True
                document.add_paragraph()  # Spacing
                continue
                
            analysis = result.get('analysis', {})
            params = result.get('params', {})
            section_title = params.get("sheet_name", report_type.replace("_", " ").title())
            
            # Add page break between sections
            if not first_section:
                document.add_page_break()
            first_section = False
            
            # Add section title
            document.add_heading(section_title, level=1)
            
            # Add insights
            if analysis.get('insights'):
                document.add_heading("Key Insights", level=2)
                
                for key, insight_text in analysis['insights'].items():
                    if key != 'status':
                        p = document.add_paragraph(style='List Bullet')
                        p.add_run(insight_text)
                        
                document.add_paragraph()  # Spacing
                
            # Add stats
            if analysis.get('summary_stats'):
                document.add_heading("Summary Statistics", level=2)
                
                # Filter out complex stats and internal data
                simple_stats = {
                    k: v for k, v in analysis['summary_stats'].items()
                    if not k.startswith('_') and 
                    not isinstance(v, (dict, list)) and
                    v is not None
                }
                
                if simple_stats:
                    for key, value in simple_stats.items():
                        p = document.add_paragraph()
                        p.add_run(f"{key.replace('_', ' ').title()}: ").bold = True
                        p.add_run(str(value))
                        
                document.add_paragraph()  # Spacing
                
            # Add data table (sample)
            data = result.get('data', [])
            columns = params.get('columns', [])
            
            if data and columns:
                document.add_heading("Data Sample", level=2)
                
                # Limit rows for sample
                max_rows = min(10, len(data))
                sample_data = data[:max_rows]
                
                # Create table
                table = document.add_table(rows=1, cols=len(columns))
                table.style = 'Table Grid'
                table.alignment = WD_TABLE_ALIGNMENT.CENTER
                
                # Add header row
                hdr_cells = table.rows[0].cells
                for i, col in enumerate(columns):
                    cell = hdr_cells[i]
                    cell.text = col.replace('.', ' ').replace('_', ' ').title()
                    cell_para = cell.paragraphs[0]
                    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    run = cell_para.runs[0]
                    run.bold = True
                    
                    # Set background color (light blue)
                    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="D0D8E8"/>')
                    cell._tc.get_or_add_tcPr().append(shading_elm)
                
                # Add data rows
                for row_dict in sample_data:
                    row_cells = table.add_row().cells
                    for i, col in enumerate(columns):
                        val = row_dict.get(col)
                        if isinstance(val, bool):
                            val = "Yes" if val else "No"
                        elif val is None:
                            val = ""
                        row_cells[i].text = str(val)
                        row_cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
                
                # Set column widths and autofit
                table.autofit = False
                for row in table.rows:
                    for cell in row.cells:
                        cell.width = Inches(1.0)
                
                # Add note if data was truncated
                if len(data) > max_rows:
                    p = document.add_paragraph(f"Note: Showing {max_rows} of {len(data)} records.")
                    p.style = document.styles['Caption']
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    
                document.add_paragraph()  # Spacing
                
            # Add charts
            if analysis.get('charts'):
                document.add_heading("Visual Analysis", level=2)
                
                for chart_key, chart_bytes in analysis['charts'].items():
                    if isinstance(chart_bytes, BytesIO):
                        try:
                            chart_bytes.seek(0)
                            
                            # Add chart title/caption
                            chart_title = chart_key.replace('_', ' ').title()
                            p = document.add_paragraph(chart_title)
                            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            p.style = document.styles['Caption']
                            
                            # Add chart image
                            document.add_picture(chart_bytes, width=Inches(6.0))
                            last_paragraph = document.paragraphs[-1]
                            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                            
                            document.add_paragraph()  # Spacing
                        except Exception as e:
                            logger.error(f"Error adding chart {chart_key} to DOCX: {e}")
                            p = document.add_paragraph(f"Error displaying chart: {chart_key}")
                            p.style.font.italic = True
        
        # Save document
        try:
            document.save(buffer)
        except Exception as e:
            logger.error(f"Error saving DOCX: {e}")
            
            # Create error document
            buffer = BytesIO()
            error_doc = docx.Document()
            error_doc.add_heading("Error Generating DOCX Report", level=0)
            error_doc.add_paragraph(str(e))
            
            try:
                error_doc.save(buffer)
            except Exception:
                return BytesIO(b"Failed to generate DOCX report.")
                
        buffer.seek(0)
        return buffer
        
# Helper function to parse XML for shading cells
def parse_xml(xml_string):
    """Parse an XML string and return an element."""
    from docx.oxml import parse_xml as _parse_xml
    from docx.oxml.ns import nsdecls
    return _parse_xml(xml_string)

# Helper for namespace declarations    
def nsdecls(*namespaces):
    """Return namespace declarations for specified namespaces."""
    from docx.oxml.ns import nsdecls as _nsdecls
    return _nsdecls(*namespaces)